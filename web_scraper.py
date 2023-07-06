import docx2txt
import glob
import io
import json
import os
import random
import requests
import subprocess

from bs4 import BeautifulSoup
from datetime import datetime
from docx import Document
from urllib.parse import urljoin

class WebScraper:
    def __init__(self, config):
        # Initialize any necessary variables or configurations
        self.config = config
        self.website = self.config['website']
        self.positions_file = self.config['positions_file']
        self.filtered_positions_file = self.config['filtered_positions_file']
        self.filter_keywords = self.config['filter_keywords']
        self.job = self.config['job']
        self.location = self.config['location']
        self.remove = self.config['remove']
        self.my_keywords_file = self.config['my_keywords_file']
        self.template_cl = self.config['template_cl']
        self.output_file_prefix = f"{self.config['output_file_prefix']}-{self.job}"
        self.cl_positions_file = self.config['cl_positions_file']
        self.open_positions = self.config['open_positions']
        self.number_of_tabs = self.config['number_of_tabs']
        self.apply_positions = self.config['apply_positions']
        self.applied_folder = self.config['applied_folder']
        self.deleted_folder = self.config['deleted_folder']

    def get_unwanted_links(self, applied_folder, deleted_folder):
        paths_list = glob.glob(os.path.join(applied_folder, '*.docx'))
        paths_list.extend(glob.glob(os.path.join(deleted_folder, '*.docx')))
        links = []

        for path in paths_list:
            text = docx2txt.process(path)
            for word in text.split():
                if word.startswith('https:'):
                    links.append(word)
                    break
        return links
    
    def write_to_file(self, positions, file_to_write):
            with open(file_to_write, 'w') as file:
                json.dump(positions, file, indent=4)

    def apply(self, positions_file):
        with open(positions_file, 'r') as f:
            positions = json.load(f)
            for position in positions:
                url = position['link'].split('?')[0] + '/apply'
                if position == positions[0]:
                    subprocess.run(['open', '-n', '-a', 'Google Chrome', '--args', '--new-window', url])
                else:
                    subprocess.run(['open', '-n', '-a', 'Google Chrome', '--args', '--new-tab', url])

    def open_links(self, positions_file):
        with open(positions_file, 'r') as f:
            positions = json.load(f)
            tabs_count = 0

            for position in positions:
                url = position['link']
                if tabs_count < self.number_of_tabs:
                    if position == positions[0]:
                        subprocess.run(['open', '-n', '-a', 'Google Chrome', '--args', '--new-window', url])
                    else:
                        subprocess.run(['open', '-n', '-a', 'Google Chrome', '--args', '--new-tab', url])
                    tabs_count += 1
                else:
                    break  # Break out of the loop if the maximum number of tabs is reached

    def create_cl(self, input_filename, output_file_prefix, positions_file, my_keywords_file):
        with open(input_filename, 'rb') as f:
            document_bytes = f.read()
        with open(positions_file, 'r') as f:
            positions = json.load(f)
        with open(my_keywords_file, 'r') as f:
            my_keys = json.load(f)
        today = datetime.today()
        formatted_date = today.strftime("%e %B %Y")

        for position in positions:
            print(f"Creating CL for {position['title']}")
            title = position['title']
            company = position['company']
            link = position['link']

            document = Document(io.BytesIO(document_bytes))

            for paragraph in document.paragraphs:
                if 'the_date' in paragraph.text:
                    paragraph.text = paragraph.text.replace('the_date', formatted_date)

                if 'the_role' in paragraph.text:
                    paragraph.text = paragraph.text.replace('the_role', title)

                if 'the_company' in paragraph.text:
                    paragraph.text = paragraph.text.replace('the_company', company) 

                if 'the_link' in paragraph.text:
                    paragraph.text = paragraph.text.replace('the_link', link)                             
            
            job_keys = position['job_keys']            
            k = 0
            for i in range(1, 5):
                if job_keys == []:
                    popped_key = my_keys[k]
                    k += 1
                else:
                    popped_key = job_keys.pop(0)
                for paragraph in document.paragraphs:
                    if f'the_key_{i}' in paragraph.text:
                        for run in paragraph.runs:
                            if f'the_key_{i}' in run.text:
                                pref = ['Experience in '+popped_key, popped_key.capitalize()+" skills"]
                                random_index = random.randrange(len(pref))
                                if len(popped_key.split()) == 1:
                                    run.text = run.text.replace(f'the_key_{i}', pref[random_index])
                                else:
                                    run.text = run.text.replace(f'the_key_{i}', popped_key[:1].capitalize() + popped_key[1:])

            index = 1
            output_filename = f'cover_letters/{output_file_prefix}-{title}.docx'
            while os.path.exists(output_filename):
                output_filename = f'cover_letters/{output_file_prefix}-{title}_{index}.docx'
                index += 1

            os.makedirs(os.path.dirname(output_filename), exist_ok=True)
            document.save(output_filename)
        
    def filter_positions(self, file, keywords):
        with open(file, 'r') as f:
            positions = json.load(f)
        new_positions = []
        for position in positions:
            print(f"Filtering {position['title']}")
            url = position['link'] 
            response = requests.get(url)
            soup = BeautifulSoup(response.content, "html.parser")   
            page_text = soup.get_text().lower()
            if len([key for key in keywords if key.lower() in page_text]) != 0:
                if self.remove and len([key for key in self.remove if key.lower() in page_text]) == 0:
                    new_positions.append(position)
        return new_positions
    
    def check_repeated_keywords(self, my_keywords_file):
        with open(my_keywords_file, 'r') as f:
            job_keywords = json.load(f) 

            repeated_keywords = [(keyword1, keyword2) for keyword1 in job_keywords for keyword2 in job_keywords 
                                 if (keyword1 in keyword2 or keyword2 in keyword1)
                                 and keyword1 != keyword2]

            if repeated_keywords:
                print("Partially repeated keywords:")
                for keyword1, keyword2 in repeated_keywords:
                    print(f"{keyword1} is partially repeated in {keyword2}")
            else:
                print("No partially repeated keywords found.")

    def find_keywords(self, positions_file, my_keywords_file):
        with open(positions_file, 'r') as f:
            positions = json.load(f)

        with open(my_keywords_file, 'r') as f:
            job_keywords = json.load(f)    

        for position in positions:
            print("Keywords for:", position['title'])
            url = position['link'] 
            response = requests.get(url)
            soup = BeautifulSoup(response.content, "html.parser")   
            page_text = soup.get_text().lower()
            job_keys = set([key for key in job_keywords if f'{key.lower()}' in page_text])
            position['job_keys'] = list(job_keys)
        return positions

    def get_seek_positions(self):
        links_to_remove = self.get_unwanted_links(self.applied_folder, self.deleted_folder)
        positions = []
        page = 1
        while True:
            key = ('-').join(self.job.split())
            url = f"{self.website}/{key}-jobs?page={page}" 
            response = requests.get(url)
            soup = BeautifulSoup(response.content, "html.parser")
            job_elements = soup.select("article[data-automation='normalJob']")
            
            if not job_elements:
                break
            for job_element in job_elements: 
                aria_label = job_element['aria-label']
                if aria_label:
                    title = aria_label
                else:
                    title = "N/A"

                company_element = job_element.select_one("a[data-automation='jobCompany']")
                if company_element:
                    company = company_element.get_text(strip=True)
                else:
                    company = "N/A"
                
                location_element = job_element.select_one("a[data-automation='jobLocation']")
                if location_element:
                    job_location = location_element.get_text(strip=True)
                else:
                    job_location = "N/A"
                
                job_link_element = job_element.select_one("a[data-automation='jobTitle']")
                if job_link_element:
                    job_link = job_link_element['href']
                    job_link = urljoin(url, job_link)
                else:
                    job_link = None
                
                # Remove special characters and unvanted symbols
                special_characters = ['ā', 'ē', 'ī', 'ō', 'ū']
                english_replacements = ['a', 'e', 'i', 'o', 'u']
                coupled_list = list(zip(special_characters,english_replacements))
                for s_char, e_char in coupled_list:
                    title = title.replace(s_char, e_char)
                    company = company.replace(s_char, e_char)
                    job_location = job_location.replace(s_char, e_char)

                unwanted_symbols = "\$#^&*@+=()/!?,[]{}~`<>,"
                for symbol in unwanted_symbols:
                    title = title.replace(symbol, "")
                    company = company.replace(symbol, "")
                    job_location = job_location.replace(symbol, "")

                if (
                    set(job_location.lower().split()) & set(self.location) 
                    and not any(word in self.remove for word in title.lower().split())
                    and job_link not in links_to_remove
                ):
                    print(f'Getting: {title}')
                    position = {
                        "title": title,
                        "company": company,
                        "location": job_location,
                        "link": job_link
                    }
                    positions.append(position)
            page += 1
        return positions
    